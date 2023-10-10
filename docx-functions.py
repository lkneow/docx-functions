import docx
from docx.enum.section import WD_ORIENT
from docx.oxml.shared import qn
from docx.oxml.xmlchemy import OxmlElement
from docx.shared import Pt
from docx.table import _Cell


def set_cell_bg_color(tc: _Cell, colorhex: str = "00519E"):
    """
    Sets the cell background shading
    
    tc: Cell object
    colorhex: hexadecimal colour codes
    """
    tblCellProperties = tc._element.tcPr
    clShading = OxmlElement('w:shd')
    clShading.set(qn('w:fill'), colorhex)  # Hex of Dark Blue Shade {R:0x00, G:0x51, B:0x9E}
    tblCellProperties.append(clShading)

def set_cell_borders(tc: _Cell, val: str = "single", sz: str = "1", color: str = "#000000"):
    '''
    Sets the cell borders
    
    tc: Cell object
    val: style of border. With reference to http://officeopenxml.com/WPtableBorders.php
    sz: width of border. With reference to http://officeopenxml.com/WPtableBorders.php
    color: color of border. In hexadecimal colour code. 
    '''
    tblCellProperties = tc._element.tcPr
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ["top", "bottom", "left", "right"]:
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), val)
        border.set(qn("w:sz"), sz)
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tblCellProperties.append(tcBorders)

def add_cell_text(tc: _Cell, text: str = "text", font:str = "Calibri", sz: int = 11, bold: bool = False):
    '''
    Adds text to table cell
    
    tc: Cell object
    text: The text to add in the cell
    font: font of the text
    sz: font size
    bold: bold or not bold
    '''
    tc.text = text
    tc.paragraphs[0].runs[0].font.name = font
    tc.paragraphs[0].runs[0].font.size = Pt(sz)
    tc.paragraphs[0].runs[0].font.bold = bold

def doc_base(font:str = "Calibri"):
    '''
    Example of initialising a document
    
    initialises document
    sets font for whole document
    sets orientation, from portrait to landscape
    '''
    doc = docx.Document()
    
    font = doc.styles['Normal'].font
    font.name = "Calibri"
    
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    return doc